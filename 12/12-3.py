import configparser
import os
import docx
from docx.document import Document
from docx.shared import Pt

os.chdir(os.path.dirname(os.path.abspath(__file__)))

conf = configparser.ConfigParser()
conf.read('conf.ini')

print(conf.sections())
conf_default = conf['DEFAULT']

doc:docx.document.Document = docx.Document(conf_default['wfile'])

doc.add_heading('가장 큰 제목', level=0)
doc.add_paragraph('여기에 원하는 텍스트를 마음껏 입력하면 됩니다.')
p = doc.add_paragraph('두번째 텍스트. 뭔가를 바꾸어 본다. Run 이용')
p.add_run('문단에 굵은 글자 추가').bold = True
p.add_run('문단에 이텔릭 글자 추가').italic = True
p.add_run('문단에 밑줄 글자 추가').underline = True
doc.add_paragraph('새문단은 자동으로 줄바꿈')

style = doc.styles["Normal"]
font = style.font
font.name = "Arial"

for p in doc.paragraphs:
    for r in p.runs:
        r.font.size = Pt(20)

doc.save(conf_default['wfile'])